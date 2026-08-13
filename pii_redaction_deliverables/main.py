from pathlib import Path
from docx import Document

from src.llm_verifier import verify_candidates
from src.redactor import redact_document
from src.detectors import (
    detect_emails,
    detect_phones,
    normalize_phone,
    detect_ips,
    detect_credit_cards,
    detect_ssns,
    detect_dobs,
    detect_names,
    detect_companies,
    detect_addresses
)


# ============================================================
# FILE PATHS
# ============================================================

INPUT_FILE = Path(
    "input/Red Herring Prospectus.docx"
)

OUTPUT_FILE = Path(
    "output/Red Herring Prospectus_REDACTED.docx"
)


# ============================================================
# EXTRACT TEXT FROM DOCX
# ============================================================

def extract_text(document) -> str:

    parts = []

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                text = cell.text.strip()

                if text:
                    parts.append(text)

    return "\n".join(parts)


# ============================================================
# MAIN
# ============================================================

def main():

    # ========================================================
    # LOAD DOCUMENT
    # ========================================================

    print("\n==========================================")
    print("LOADING DOCUMENT")
    print("==========================================")

    document = Document(INPUT_FILE)

    text = extract_text(document)

    print("Input file:", INPUT_FILE)
    print("Total extracted characters:", len(text))


    # ========================================================
    # EMAIL
    # ========================================================

    print("\n==========================================")
    print("EMAIL DETECTION")
    print("==========================================")

    emails = detect_emails(text)

    print("Email occurrences:", len(emails))
    print(
        "Unique email addresses:",
        len(set(x["text"] for x in emails))
    )


    # ========================================================
    # PHONE
    # ========================================================

    print("\n==========================================")
    print("PHONE DETECTION")
    print("==========================================")

    phones = detect_phones(text)

    print("Phone occurrences:", len(phones))
    print(
        "Unique phone numbers:",
        len(set(x["text"] for x in phones))
    )
    print(
        "Unique normalized phones:",
        len(set(
            normalize_phone(x["text"])
            for x in phones
        ))
    )


    # ========================================================
    # IP ADDRESS
    # ========================================================

    print("\n==========================================")
    print("IP ADDRESS DETECTION")
    print("==========================================")

    ips = detect_ips(text)

    print("IP occurrences:", len(ips))
    print(
        "Unique IP addresses:",
        len(set(x["text"] for x in ips))
    )


    # ========================================================
    # CREDIT CARD
    # ========================================================

    print("\n==========================================")
    print("CREDIT CARD DETECTION")
    print("==========================================")

    credit_cards = detect_credit_cards(text)

    print(
        "Credit card occurrences:",
        len(credit_cards)
    )

    print(
        "Unique credit cards:",
        len(set(x["text"] for x in credit_cards))
    )


    # ========================================================
    # SSN
    # ========================================================

    print("\n==========================================")
    print("SSN DETECTION")
    print("==========================================")

    ssns = detect_ssns(text)

    print("SSN occurrences:", len(ssns))
    print(
        "Unique SSNs:",
        len(set(x["text"] for x in ssns))
    )


    # ========================================================
    # DOB
    # ========================================================

    print("\n==========================================")
    print("DATE OF BIRTH DETECTION")
    print("==========================================")

    dobs = detect_dobs(text)

    print("DOB occurrences:", len(dobs))
    print(
        "Unique DOBs:",
        len(set(x["text"] for x in dobs))
    )

    
    # ========================================================
    # ADDRESS
    # ========================================================

    print("\n==========================================")
    print("ADDRESS DETECTION")
    print("==========================================")

    addresses = detect_addresses(text)

    unique_addresses = set(
        address["text"]
        for address in addresses
    )

    print(
        "Address occurrences:",
        len(addresses)
    )

    print(
        "Unique addresses:",
        len(unique_addresses)
    )

    # ========================================================
    # NAME
    # ========================================================

    print("\n==========================================")
    print("NAME DETECTION")
    print("==========================================")

    names = detect_names(text)

    print(
        "Name occurrences:",
        len(names)
    )

    print(
        "Unique names:",
        len(set(
            x["text"]
            for x in names
        ))
    )


    # ========================================================
    # COMPANY
    # ========================================================

    print("\n==========================================")
    print("COMPANY DETECTION")
    print("==========================================")

    companies = detect_companies(text)

    print(
        "Company occurrences:",
        len(companies)
    )

    print(
        "Unique companies:",
        len(set(
            x["text"]
            for x in companies
        ))
    )


    # ========================================================
    # COMBINE RULE-BASED DETECTIONS
    # ========================================================

    print("\n==========================================")
    print("COMBINING DETECTIONS")
    print("==========================================")

    all_detections = (
        emails
        + phones
        + ips
        + credit_cards
        + ssns
        + dobs
        + names
        + companies
        + addresses

    )

    print(
        "TOTAL RULE-BASED CANDIDATES:",
        len(all_detections)
    )


    # ========================================================
    # GROQ NAME VERIFICATION
    # ========================================================

    print("\n==========================================")
    print("GROQ SECONDARY VERIFICATION")
    print("==========================================")

    print(
        "Sending NAME candidates to Groq..."
    )

    verified_candidates = verify_candidates(
        text,
        all_detections
    )

    print(
        "\nVerified names:",
        len(verified_candidates)
    )


    # ========================================================
    # PRINT VERIFIED NAMES
    # ========================================================

    print("\n==========================================")
    print("FINAL VERIFIED PERSON NAMES")
    print("==========================================")

    for candidate in verified_candidates:

        print(
            f'{candidate["text"]} | '
            f'confidence={candidate["confidence"]} | '
            f'start={candidate["start"]} | '
            f'end={candidate["end"]}'
        )


    # ========================================================
    # BUILD FINAL REDACTION LIST
    # ========================================================

    print("\n==========================================")
    print("BUILDING FINAL REDACTION LIST")
    print("==========================================")

    # Everything except ambiguous names is
    # deterministically redacted.

    deterministic_pii = (
        emails
        + phones
        + ips
        + credit_cards
        + ssns
        + dobs
        + companies
        + addresses
    )

    final_pii = (
        deterministic_pii
        + verified_candidates
    )

    print(
        "Deterministic PII:",
        len(deterministic_pii)
    )

    print(
        "LLM verified names:",
        len(verified_candidates)
    )

    print(
        "Total redaction candidates:",
        len(final_pii)
    )


    # ========================================================
    # REMOVE DUPLICATES
    # ========================================================

    unique_pii = {}

    for pii in final_pii:

        key = (
            pii["start"],
            pii["end"],
            pii["text"]
        )

        unique_pii[key] = pii

    final_pii = list(
        unique_pii.values()
    )

    print(
        "Unique redaction candidates:",
        len(final_pii)
    )


    # ========================================================
    # CREATE OUTPUT DIRECTORY
    # ========================================================

    OUTPUT_FILE.parent.mkdir(
        parents=True,
        exist_ok=True
    )


    # ========================================================
    # REDACT
    # ========================================================

    print("\n==========================================")
    print("REDACTING DOCUMENT")
    print("==========================================")

    redact_document(
        INPUT_FILE,
        OUTPUT_FILE,
        final_pii
    )


    # ========================================================
    # FINAL SUMMARY
    # ========================================================

    print("\n==========================================")
    print("FINAL SUMMARY")
    print("==========================================")

    print(
        "Total characters:",
        len(text)
    )

    print(
        "Emails:",
        len(emails)
    )

    print(
        "Phones:",
        len(phones)
    )

    print(
        "IPs:",
        len(ips)
    )

    print(
        "Credit cards:",
        len(credit_cards)
    )

    print(
        "SSNs:",
        len(ssns)
    )

    print(
        "DOBs:",
        len(dobs)
    )

    print(
        "Addresses:",
        len(addresses)
    )

    print(
        "Names detected:",
        len(names)
    )

    print(
        "Companies detected:",
        len(companies)
    )

    print(
        "Rule-based candidates:",
        len(all_detections)
    )

    print(
        "LLM verified names:",
        len(verified_candidates)
    )

    print(
        "Final redaction candidates:",
        len(final_pii)
    )

    print(
        "\nRedacted document:"
    )

    print(
        OUTPUT_FILE
    )

    print("==========================================")


# ============================================================
# PROGRAM ENTRY POINT
# ============================================================

if __name__ == "__main__":
    main()