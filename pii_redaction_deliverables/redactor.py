from docx import Document


REDACTION_TEXT = "[REDACTED]"


def redact_text_by_matches(text, pii_list):
    """
    Redact PII from a single piece of text.

    pii_list contains absolute document positions.
    We convert them to local positions before replacing.
    """

    if not pii_list:
        return text

    # Find actual occurrences of the PII text inside this text.
    matches = []

    for pii in pii_list:

        pii_text = pii["text"]

        if not pii_text:
            continue

        start = text.find(pii_text)

        if start == -1:
            continue

        end = start + len(pii_text)

        matches.append({
            "start": start,
            "end": end
        })

    # Remove overlapping matches
    matches.sort(
        key=lambda x: (x["start"], x["end"]),
        reverse=True
    )

    redacted = text

    for match in matches:

        start = match["start"]
        end = match["end"]

        redacted = (
            redacted[:start]
            + REDACTION_TEXT
            + redacted[end:]
        )

    return redacted


def redact_document(input_file, output_file, verified_pii):
    """
    Create a redacted DOCX.

    Redacts verified PII from:
    - paragraphs
    - table cells
    """

    document = Document(input_file)

    total_redactions = 0

    # ==========================================
    # PARAGRAPHS
    # ==========================================

    for paragraph in document.paragraphs:

        original_text = paragraph.text

        if not original_text:
            continue

        paragraph_pii = []

        for pii in verified_pii:

            if pii["text"] in original_text:
                paragraph_pii.append(pii)

        if not paragraph_pii:
            continue

        new_text = redact_text_by_matches(
            original_text,
            paragraph_pii
        )

        if new_text != original_text:
            total_redactions += len(paragraph_pii)
            paragraph.text = new_text

    # ==========================================
    # TABLES
    # ==========================================

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    original_text = paragraph.text

                    if not original_text:
                        continue

                    cell_pii = []

                    for pii in verified_pii:

                        if pii["text"] in original_text:
                            cell_pii.append(pii)

                    if not cell_pii:
                        continue

                    new_text = redact_text_by_matches(
                        original_text,
                        cell_pii
                    )

                    if new_text != original_text:
                        total_redactions += len(cell_pii)
                        paragraph.text = new_text

    # ==========================================
    # SAVE
    # ==========================================

    output_file.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    document.save(output_file)

    print("\n==========================================")
    print("REDACTION COMPLETE")
    print("==========================================")
    print("Redactions applied:", total_redactions)
    print("Output file:", output_file)
    print("==========================================")